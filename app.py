import os
import uuid
import json
import subprocess
from io import BytesIO

from dotenv import load_dotenv

from flask import Flask, request, flash, redirect, render_template, send_file, session, abort, after_this_request

from docx import Document

from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

app = Flask(__name__)

load_dotenv()
libre_office = os.getenv("LIBRE_OFFICE_PATH")

ALLOWED_EXTENSIONS = {'docx'}
UPLOAD_FOLDER   = os.path.join(os.getcwd(), "uploads")
TEMPLATE_FOLDER = os.path.join(os.getcwd(), "docx_templates")

app.secret_key = "dev-secret-key"
app.config['MAX_CONTENT_LENGTH'] = 10 * 1000 * 1000 #10 MB
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMPLATE_FOLDER'], exist_ok=True)

# Might need to save the document temporarily to keep it in memory or something so it can be altered later

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/", methods=['GET', 'POST'])
def init_page():

    wipe_uploads_folder()

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'upload':
            try:
                if 'file' in request.files:
                    # Check if file actually submitted
                    # Get file
                    file = request.files['file']

                    if file.filename == '':
                        flash('No file selected')
                        return redirect(request.url)
                    
                    # Check file exists and is allowed
                    if file and allowed_file(file.filename):
                        source_stream = BytesIO(file.read())
                        source_stream.close()

                        # Save uploaded file temporarily and the path in session
                        fName = (file.filename)
                        path = os.path.join(app.config['UPLOAD_FOLDER'], fName)
                        file.stream.seek(0)
                        file.save(path)
                        session['selected_template'] = fName
                        session['selected_template_location'] = app.config['UPLOAD_FOLDER']

                        return redirect('/create-pdf')

                    else:
                        flash('Accepted filetypes: ' + str(ALLOWED_EXTENSIONS).strip('}{'))
                        return redirect(request.url)
                else:
                    flash('File not found in request')
                    return redirect(request.url)         

            except RequestEntityTooLarge:
                flash('Maximum filesize: ' + str(app.config['MAX_CONTENT_LENGTH'] / (1000 * 1000)) + ' MB' )
                return redirect(request.url)

        if action == 'docSelect':
            filename = request.form['docSelect']
            if filename:
                session['selected_template'] = filename
                session['selected_template_location'] = app.config['TEMPLATE_FOLDER']
                return redirect('/create-pdf')  
        
    path = app.config['TEMPLATE_FOLDER']
    filenames = os.listdir(path)
    return render_template("template_hub.html", fNames=filenames, allowed_extensions=ALLOWED_EXTENSIONS)

@app.route('/create-pdf', methods=['GET', 'POST'])
def create_pdf_page():
    if request.method == 'POST':
        action = request.form.get('action')

        ## Handle if download button is clicked ## 
        if action == 'download':
            # All submitted fields
            form_data = request.form.to_dict()

            # Remove non-field keys (like the document title etc.)
            form_data.pop('action', None)
            form_data.pop('download', None)

            return handle_download(form_data)

    ## Create the download from template page ##    
    
    filename = session.get('selected_template')
    filePath = session.get('selected_template_location')

    if not filename or not filePath:
        abort(400, "No template selected")

    path = os.path.join(filePath, filename)

    if not os.path.isfile(path):
        abort(404, "Template not found")

    document = Document(path)

    full_text = [p.text for p in document.paragraphs]
    fieldArr = []

    for paragraph in document.paragraphs:
        text = paragraph.text
        start = text.find("[--")

        while start != -1:
            end = text.find("--]", start)
            if end == -1:
                break

            fieldText = text[start:end + 3]

            if fieldText not in fieldArr:
                fieldArr.append(text[start:end + 3])  # Include the "--]"

            start = text.find("[--", end + 3)
    
    return render_template('create_new_pdf.html', text=full_text, fileName=filename, fields=fieldArr)

@app.route('/create-template', methods=['GET', 'POST'])
def create_template_page():
    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'upload':
            return handle_upload(request)         

        elif action == 'selections':
            return handle_selections(request)
            
        elif action == 'fields':
            return handle_fields(request)

        else:
            abort(400, "Invalid or missing action")
            
    else:
        return render_template('create_template.html', allowed_extensions=ALLOWED_EXTENSIONS)
                
def get_current_doc_path():
    doc_path = session.get('doc_path')

    if not doc_path:
        abort(400)

    if not os.path.isfile(doc_path):
        app.logger.error(f"Missing docx: {doc_path}")
        flash('Document file does not exist, try a different file')
        return redirect(request.url)

    if os.path.getsize(doc_path) == 0:
        flash('Document file is empty, try a different file')
        return redirect(request.url)
    
    return doc_path

def get_current_selected_template():

    filename = session.get('selected_template')
    filePath = session.get('selected_template_location')

    if not filename or not filePath:
        abort(400, "No template selected")

    path = os.path.join(filePath, filename)

    if not os.path.isfile(path):
        app.logger.error(f"Missing docx: {path}")
        flash('Document file does not exist, try a different file')
        return redirect(request.url)

    if os.path.getsize(path) == 0:
        flash('Document file is empty, try a different file')
        return redirect(request.url)

    return path

def handle_upload(request):
    # Check if the post request has the file part
            try:
                if 'file' in request.files:
                    # Check if file actually submitted
                    # Get file
                    file = request.files['file']

                    if file.filename == '':
                        flash('No file selected')
                        return redirect(request.url)
                    
                    # Check file exists and is allowed
                    if file and allowed_file(file.filename):
                        source_stream = BytesIO(file.read())
                        document = Document(source_stream)
                        source_stream.close()

                        # Save uploaded file temporarily and the path in session
                        fName = secure_filename(file.filename)
                        path = os.path.join(app.config['UPLOAD_FOLDER'], fName)
                        file.stream.seek(0)
                        file.save(path)
                        session['doc_path'] = path

                        # Extract text from all paragraphs
                        full_text = []
                        for paragraph in document.paragraphs:
                            full_text.append(paragraph.text)

                        return render_template('create_template.html', text=full_text, fileName=file.filename, allowed_extensions=ALLOWED_EXTENSIONS)
                    else:
                        flash('Accepted filetypes: ' + str(ALLOWED_EXTENSIONS).strip('}{'))
                        return redirect(request.url)
                else:
                    flash('File not found in request')
                    return redirect(request.url)         

            except RequestEntityTooLarge:
                flash('Maximum filesize: ' + str(app.config['MAX_CONTENT_LENGTH'] / (1000 * 1000)) + ' MB' )
                return redirect(request.url)
            
def handle_selections(request):
    try:
        selections = json.loads(request.form['selections'])
        # return f"<div>{selections}</div>"

        doc_path = get_current_doc_path()    
        document = Document(doc_path)

        # We're going to march through this char by char and replace exactly our selection, keeping all that formatting crap intact
        # We, in fact, did not do that
        full_text    = []
        final_fields = []

        for pIndex, paragraph in enumerate(document.paragraphs, start=1):

            preview_text    = ""
            reference_text  = ""
            temp_selections = []
            
            for run in paragraph.runs:
                reference_text += run.text
                preview_text   += run.text

            for selection in selections:                
                selText        = "".join(ch for ch in selection['text'] if not ch.isspace()).lower()
                pHolderText    = selection['placeholder']

                stripped_text  = []
                original_index = []

                for index, char in enumerate(reference_text):
                    if not char.isspace():
                        stripped_text.append(char)
                        original_index.append(index)

                stripped_text = "".join(stripped_text).lower()

                # Gotta find all matches within paragraph text
                start_stripped = stripped_text.find(selText)
                while start_stripped != -1:
                    start_original = original_index[start_stripped]
                    end_original   = original_index[start_stripped + len(selText) - 1]

                    temp_selection = {
                        "uuid": str(uuid.uuid4()),
                        "text": selection['text'],
                        "paragraphIndex": pIndex,
                        "index": {"start": start_original, "end": end_original},
                        "field_text": pHolderText
                    }
                    temp_selections.append(temp_selection)
                    final_fields.append(temp_selection)

                    # Replace the matched portion in stripped_text to avoid finding the same one again
                    stripped_text = stripped_text[:start_stripped] + (" " * len(selText)) + stripped_text[start_stripped + len(selText):]
                    start_stripped = stripped_text.find(selText)

            temp_selections.sort(key=lambda sel: sel["index"]["start"], reverse=True)

            for sel in temp_selections:
                start_original = sel['index']['start']
                end_original   = sel['index']['end']
                pHolderText    = sel['field_text']

                # Create new string by combining slices and a new character
                preview_text = preview_text[:start_original] + pHolderText + preview_text[end_original + 1:]
                    
            full_text.append(preview_text)

        return render_template('save_template.html', text=full_text, fields=final_fields, title_placeholder=str(uuid.uuid4()))
    
    except KeyError:
        # Missing 'selections' in POST
        flash("No selection data received from the form.")
        return redirect(request.url)

    except Exception as e:
        # Catch everything else lol
        app.logger.error(f"Unexpected error in handle_selections: {e}")
        flash("An unexpected error occurred. Please try again.")
        return redirect(request.url)
    
def handle_fields(request):
    try:
        fields = json.loads(request.form['fields'])
        template_name = fields.pop()['templateName']
        # return f"<div>{fields}</div>"

        doc_path = get_current_doc_path()    
        document = Document(doc_path)
        
        for pIndex, paragraph in enumerate(document.paragraphs, start=1):
            # Sort fields in reverse based on where they are in paragraph
            para_fields = [f for f in fields if f["paragraphIndex"] == pIndex]
            para_fields.sort(key=lambda f: f["index"]["start"], reverse=True)
            flat = ""
            runs = []
            # Combine runs into flat text, map the start and end indexes per paragraph 
            for run in paragraph.runs:
                start = len(flat)
                flat += run.text
                end = len(flat)
                runs.append({
                    "run": run,
                    "start": start,
                    "end": end
                })
                
            # Using indexes from the flattened list and the field indexes
            for field in para_fields:
                if field['paragraphIndex'] == pIndex:
                    field_start = field["index"]["start"]
                    field_end   = field["index"]["end"] + 1
                    replacement = field["field_text"]

                    inserted = False

                    for r in runs:
                        if r["end"] <= field_start or r["start"] >= field_end:
                            continue

                        run = r["run"]
                        text = run.text

                        local_start = max(field_start, r["start"]) - r["start"]
                        local_end   = min(field_end, r["end"]) - r["start"]

                        if not inserted:
                            run.text = text[:local_start] + replacement + text[local_end:]
                            inserted = True
                        else:
                            run.text = text[:local_start] + text[local_end:]
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)

        fName = (template_name + ".docx")
        upPath = os.path.join(app.config['UPLOAD_FOLDER'], fName)
        document.save(upPath)        

        if request.form.get('sub-action') == "save-local":
            # Load that bad boy into memory
            file_stream = BytesIO()
            with open(upPath, "rb") as f:
                file_stream.write(f.read())
            file_stream.seek(0)

            # Download it
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=fName,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        if request.form.get('sub-action') == "create":
            file_stream = BytesIO()
            with open(upPath, "rb") as f:
                file_stream.write(f.read())
            path = os.path.join(app.config['TEMPLATE_FOLDER'], fName)
            document.save(path)            

        # Save the new template to the template folder (for now, upgrade to db storage in future potentially)
        # fName = secure_filename(template_name + ".docx")
        # The werkzeug secure_filename function was removing char I wanted so we're rawdogging it now

        return redirect('/')  


    except KeyError:
        # Missing 'selections' in POST
        flash("No selection data received from the form.")
        return redirect(request.url)

    except Exception as e:
        # Catch everything else lol
        app.logger.error(f"Unexpected error in handle_fields: {e}")
        flash("An unexpected error occurred. Please try again.")
        return redirect(request.url)

def handle_download(field_values):
    selected_template = get_current_selected_template()
    document = Document(selected_template)
    filename = field_values['docName'].strip()

    for key, value in field_values.items():
        if key in filename:
            fValue = "".join(str(value).split())
            filename = filename.replace(str(key), fValue)

    # To note: if the value to replace has a '.' in it, the rest of the filename will be lobbed off in the next loop

    endFname = filename.find('.')
    if endFname != -1:
        if len(filename[:endFname].strip()) == 0:
            filename = "newfile.pdf"
        else:
            filename = filename[:endFname] + ".pdf"
    else:
        filename += ".pdf"

    for paragraph in document.paragraphs:
        for key, value in field_values.items():
            replace_placeholder_in_paragraph(paragraph, key, value.strip())

    temp_docx = save_docx_temp(document)
    pdf_file  = convert_docx_to_pdf(temp_docx)

    # Clean up DOCX
    try:
        os.remove(temp_docx)
    except Exception as e:
        app.logger.error(f"Failed to remove temp DOCX: {e}")

    pdf_bytes = BytesIO()
    with open(pdf_file, "rb") as f:
        pdf_bytes.write(f.read())
    pdf_bytes.seek(0)

    # Clean up PDF
    try:
        os.remove(pdf_file)
    except Exception as e:
        app.logger.error(f"Failed to remove temp PDF: {e}")

    return send_file(
        pdf_bytes,
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf"
    )

def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    while True:
        flat = ""
        runs = []

        # rebuild flattened text + run map
        for run in paragraph.runs:
            start = len(flat)
            flat += run.text
            end = len(flat)
            runs.append({
                "run": run,
                "start": start,
                "end": end
            })

        start = flat.find(placeholder)
        if start == -1:
            break

        end = start + len(placeholder)
        inserted = False

        for r in runs:
            if r["end"] <= start or r["start"] >= end:
                continue

            run = r["run"]
            text = run.text

            local_start = max(start, r["start"]) - r["start"]
            local_end   = min(end, r["end"]) - r["start"]

            if not inserted:
                run.text = text[:local_start] + replacement + text[local_end:]
                inserted = True
            else:
                run.text = text[:local_start] + text[local_end:]


def save_docx_temp(document):
    fName = secure_filename(str(uuid.uuid4()) + ".docx")
    path = os.path.join(app.config['UPLOAD_FOLDER'], fName)

    # Save the document to the filesystem
    document.save(path)

    # Return the full path for later use
    return path

def convert_docx_to_pdf(docx_path, output_dir=None):
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)

    # Call LibreOffice in headless mode
    subprocess.run([
        libre_office,
        "--headless",
        "--convert-to", "pdf",
        docx_path,
        "--outdir", output_dir
    ], check=True)

    # Construct PDF path
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
    return pdf_path

def wipe_uploads_folder():
    for filename in os.listdir(app.config['UPLOAD_FOLDER']): 
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename) 
        try: 
            if os.path.isfile(file_path): 
                os.remove(file_path) 
        except Exception as e:
            app.logger.error(f"Failed to delete {file_path}: {e}")


if __name__ == "__main__":
    app.run(debug=True)